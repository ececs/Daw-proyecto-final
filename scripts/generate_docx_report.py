import os
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def apply_table_borders(table):
    # Set borders for professional look (thin gray borders)
    tblPr = table._tbl.tblPr
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))

def add_page_number(run):
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_table_of_contents(paragraph):
    p = paragraph
    run = p.add_run()
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-2" \h \z \u </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_formatted_text(paragraph, text):
    # Split text by bold (**bold**) and inline code (`code`)
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11.0)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11.0)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11.0)

def generate_report():
    input_path = "MEMORIA_PROYECTO.md"
    output_path = "Memoria_D4_Ticket_AI_4.docx"

    if not os.path.exists(input_path):
        print(f"Error: {input_path} not found.")
        return

    doc = docx.Document()

    # Define standard page setup according to guidelines
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        # Enable different first page for Portada
        section.different_first_page_header_footer = True

        # Configure footer with page numbers
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.text = ""
        run_text = p.add_run("Página ")
        run_text.font.name = 'Times New Roman'
        run_text.font.size = Pt(10)
        run_num = p.add_run()
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(10)
        add_page_number(run_num)

    # Adjust Normal Style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(11.0)
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    with open(input_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # Variables for state tracking
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    # Title & Subtitle logic (Portada)
    doc.add_paragraph() # Spacer
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("D4-Ticket AI\n")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(26)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Plataforma Inteligente de Gestión de Incidencias mediante Agentes de Inteligencia Artificial\n\nInforme Final del Proyecto\n")
    sub_run.font.name = 'Times New Roman'
    sub_run.font.size = Pt(16)
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Cover Page Background Image
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_run = img_p.add_run()
    img_run.add_picture("docs_assets/d4_ticket_cover_background.png", width=Inches(4.0))

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(
        "\n"
        "Alumno: Eudaldo Álvaro Cal Saül\n"
        "Módulo profesional: Proyecto Intermodular — Ciclo Superior en Desarrollo de Aplicaciones Web (DAW)\n"
        "Profesor responsable: Óliver Díaz Rodríguez\n"
        "Fecha de entrega: Mayo de 2026\n"
    )
    info_run.font.name = 'Times New Roman'
    info_run.font.size = Pt(12)
    info_run.bold = True
    info_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    doc.add_page_break()

    # Parse body lines starting after the header metadata (skip first 16 lines of Markdown metadata)
    idx = 16
    while idx < len(lines):
        line = lines[idx].rstrip('\n')

        # Check for Índice section to replace with native Table of Contents
        if line.strip() == "## Índice":
            p = doc.add_paragraph(style='Heading 2')
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run("Índice")
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

            # Add dynamic Table of Contents field paragraph
            toc_p = doc.add_paragraph()
            add_table_of_contents(toc_p)

            doc.add_page_break()

            # Skip manual list items until next section
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("---") and not lines[idx].strip().startswith("#"):
                idx += 1
            if idx < len(lines) and lines[idx].strip().startswith("---"):
                idx += 1
            continue

        # Check for code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                # Write code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p_format = p.paragraph_format
                p_format.line_spacing = 1.0
                
                code_text = "\n".join(code_lines)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
                
                # Apply light shading to the paragraph
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F8F8"/>')
                pPr.append(pBdr)

                code_lines = []
            else:
                in_code_block = True
            idx += 1
            continue

        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue

        # Check for tables
        if line.strip().startswith("|") and not in_code_block:
            in_table = True
            table_rows.append(line)
            idx += 1
            continue
        elif in_table:
            # End of table detected or non-table line
            in_table = False
            # Parse and render table
            # Exclude header separator row (e.g. |---|---|)
            clean_rows = []
            for r in table_rows:
                if re.search(r'^[|\s:-]+$', r):
                    continue
                clean_rows.append(r)
            
            if clean_rows:
                grid = []
                for r in clean_rows:
                    cols = [c.strip() for c in r.split("|")[1:-1]]
                    grid.append(cols)

                if grid:
                    num_rows = len(grid)
                    num_cols = len(grid[0])
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.autofit = True
                    apply_table_borders(table)

                    for r_idx, row_data in enumerate(grid):
                        row = table.rows[r_idx]
                        is_header = (r_idx == 0)
                        for c_idx, val in enumerate(row_data):
                            if c_idx < len(row.cells):
                                cell = row.cells[c_idx]
                                cell.text = ""
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                p.paragraph_format.space_before = Pt(4)
                                p.paragraph_format.space_after = Pt(4)
                                run_p = p.add_run()
                                add_formatted_text(p, val)
                                if is_header:
                                    set_cell_background(cell, "F2F2F2")
                                    for run in p.runs:
                                        run.bold = True

                    # Add empty spacing paragraph after table
                    space_p = doc.add_paragraph()
                    space_p.paragraph_format.space_after = Pt(12)

            table_rows = []
            # Do not skip current line, process it as normal text below

        # Headings parsing
        if line.startswith("# "):
            p = doc.add_paragraph(style='Heading 1')
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
            idx += 1
            continue
        elif line.startswith("## "):
            p = doc.add_paragraph(style='Heading 2')
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            idx += 1
            continue
        elif line.startswith("### "):
            p = doc.add_paragraph(style='Heading 3')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            idx += 1
            continue
        elif line.startswith("#### "):
            p = doc.add_paragraph(style='Heading 4')
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[5:].strip())
            run.bold = True
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11.0)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            idx += 1
            continue

        # Bullet lists parsing
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p_format = p.paragraph_format
            p_format.line_spacing = 1.15
            content = line.strip()[2:]
            add_formatted_text(p, content)
            idx += 1
            continue

        # Blank lines
        if not line.strip():
            idx += 1
            continue

        # Image parsing: ![caption](image_path)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', line.strip())
        if img_match:
            caption_text = img_match.group(1)
            img_path = img_match.group(2)
            if os.path.exists(img_path):
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_p.add_run()
                img_run.add_picture(img_path, width=Inches(5.5))

                # Add a beautifully styled caption underneath the image
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_before = Pt(4)
                cap_p.paragraph_format.space_after = Pt(12)
                cap_run = cap_p.add_run(f"Ilustración: {caption_text}")
                cap_run.font.name = 'Times New Roman'
                cap_run.font.size = Pt(9.5)
                cap_run.italic = True
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            idx += 1
            continue

        # Normal text paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        p_format = p.paragraph_format
        p_format.line_spacing = 1.5  # Crucial 1.5 Line Spacing requirement
        add_formatted_text(p, line)
        idx += 1

    # Save document
    doc.save(output_path)
    print(f"Success: Generated {output_path} formatted exactly according to guidelines.")

if __name__ == "__main__":
    generate_report()
